import datetime
import os
import sys
from pathlib import Path

import anthropic
import requests
from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv

load_dotenv()

# --- CONFIGURATION ---
VAULT_PATH: str = os.environ.get("VAULT_PATH", "")
SCREENPIPE_URL: str = os.environ.get("SCREENPIPE_URL", "http://localhost:3030/search")
SCREENPIPE_API_KEY: str = os.environ.get("SCREENPIPE_API_KEY", "")

# Choose your AI backend: "ollama", "openai", or "claude"
AI_PROVIDER: str = os.environ.get("AI_PROVIDER", "ollama")
OLLAMA_MODEL: str = os.environ.get("OLLAMA_MODEL", "llama3")
OPENAI_API_KEY: str = os.environ.get("OPENAI_API_KEY", "")
ANTHROPIC_API_KEY: str = os.environ.get("ANTHROPIC_API_KEY", "")
CLAUDE_MODEL: str = os.environ.get("CLAUDE_MODEL", "claude-haiku-4-5")


def get_daily_text() -> str:
    """Queries Screenpipe's local API for all text & audio recorded today."""
    today_start = datetime.datetime.now(datetime.timezone.utc).replace(hour=0, minute=0, second=0, microsecond=0).strftime("%Y-%m-%dT%H:%M:%SZ")
    params = {
        "content_type": "all",
        "start_time": today_start,
        "limit": 1000,
    }

    headers = {"Authorization": f"Bearer {SCREENPIPE_API_KEY}"} if SCREENPIPE_API_KEY else {}
    try:
        response = requests.get(SCREENPIPE_URL, params=params, headers=headers, timeout=10)
        response.raise_for_status()
        data = response.json()

        text_segments: list[str] = []
        for item in data.get("data", []):
            content = item.get("content", {})
            text = content.get("text") or content.get("transcription")
            if text:
                text_segments.append(text.strip())

        return "\n".join(text_segments)
    except Exception as e:
        print(f"Error fetching from Screenpipe API: {e}")
        return ""


def summarize_text(raw_text: str) -> str:
    """Sends the daily text dump to your chosen AI for clean summarization."""
    if not raw_text.strip():
        return "No activity or text recorded on screen today."

    prompt = (
        "You are a personal knowledge manager. Analyze the following raw screen OCR and audio "
        "transcript logs from my computer today.\n"
        "Provide a clean markdown summary focusing on what I watched (like YouTube videos) or read.\n"
        "Group by major topics or apps. Keep it structural, bulleted, and actionable for long-term learning.\n\n"
        f"Raw Daily Log Data:\n{raw_text[:15000]}"
    )

    return ask_ai(prompt) or "No summary generated."


def extract_open_loops(raw_text: str, date: str) -> list[str]:
    """Second AI pass: extract unfinished tasks/intentions from today's OCR dump."""
    if not raw_text.strip():
        return []

    prompt = (
        "You are a ruthless task extractor. Look at this screen OCR / audio transcript from today.\n"
        "Extract ONLY items that are clearly unfinished, half-started, or mentioned as something "
        "the user intended to do but probably didn't complete today.\n"
        "Rules:\n"
        "- Output ONLY a plain bullet list, one item per line, starting with '- '\n"
        "- Each item max 120 characters\n"
        "- If nothing qualifies, output exactly: NONE\n"
        "- Do NOT include routine tasks, things clearly completed, or vague thoughts\n"
        "- Focus on: projects mentioned but not acted on, tabs opened but not read, "
        "things said aloud as a goal, files opened but not edited\n\n"
        f"Today's transcript:\n{raw_text[:12000]}"
    )

    raw = ask_ai(prompt)
    if not raw or raw.strip().upper() == "NONE":
        return []

    items: list[str] = []
    for line in raw.splitlines():
        line = line.strip()
        if line.startswith("- ") and len(line) > 3:
            items.append(line)
    return items


def append_to_open_loops(items: list[str], date: str) -> None:
    """Appends extracted open loops to Memory/Open-Loops.md with a date header."""
    if not items:
        return

    vault = Path(VAULT_PATH)
    open_loops = vault / "Memory" / "Open-Loops.md"
    open_loops.parent.mkdir(parents=True, exist_ok=True)

    block = f"\n## {date}\n\n" + "\n".join(items) + "\n"

    existing = open_loops.read_text(encoding="utf-8") if open_loops.exists() else ""
    if f"## {date}" in existing:
        # Replace the existing block for today (idempotent re-runs)
        lines = existing.split(f"## {date}")
        before = lines[0]
        after_parts = lines[1].split("\n## ", 1)
        rest = ("\n## " + after_parts[1]) if len(after_parts) > 1 else ""
        open_loops.write_text(before + block + rest, encoding="utf-8")
    else:
        with open_loops.open("a", encoding="utf-8") as f:
            f.write(block)

    print(f"Appended {len(items)} open loops to {open_loops}")


def ask_ai(prompt: str) -> str:
    """Shared AI call used by both summarize and extract paths."""
    if AI_PROVIDER == "ollama":
        try:
            res = requests.post(
                "http://localhost:11434/api/generate",
                json={"model": OLLAMA_MODEL, "prompt": prompt, "stream": False},
                timeout=120,
            )
            return res.json().get("response", "").strip()
        except Exception as e:
            return f"Ollama error: {e}"

    elif AI_PROVIDER == "openai":
        if not OPENAI_API_KEY:
            return "OPENAI_API_KEY not set in .env"
        try:
            res = requests.post(
                "https://api.openai.com/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {OPENAI_API_KEY}",
                    "Content-Type": "application/json",
                },
                json={"model": "gpt-4o-mini", "messages": [{"role": "user", "content": prompt}]},
                timeout=30,
            )
            return res.json()["choices"][0]["message"]["content"].strip()
        except Exception as e:
            return f"OpenAI error: {e}"

    elif AI_PROVIDER == "claude":
        if not ANTHROPIC_API_KEY:
            return "ANTHROPIC_API_KEY not set in .env"
        try:
            client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
            message = client.messages.create(
                model=CLAUDE_MODEL,
                max_tokens=2048,
                messages=[{"role": "user", "content": prompt}],
            )
            return message.content[0].text.strip()
        except Exception as e:
            return f"Claude API error: {e}"

    return f"Unknown AI_PROVIDER: {AI_PROVIDER!r}"


def send_to_obsidian(content: str) -> None:
    """Writes the formatted markdown daily note directly into the Obsidian vault folder."""
    if not VAULT_PATH:
        print("Error: VAULT_PATH is not set in .env")
        sys.exit(1)

    today = datetime.date.today().isoformat()
    folder = Path(VAULT_PATH) / "Daily_Summaries"
    folder.mkdir(parents=True, exist_ok=True)

    note = folder / f"{today}.md"
    note.write_text(
        f"# Daily Learning Summary — {today}\n\n{content}",
        encoding="utf-8",
    )
    print(f"Saved: {note}")


def send_to_word(content: str) -> None:
    """Writes the daily summary as a Word .docx file into Word_Summaries/."""
    if not VAULT_PATH:
        print("Error: VAULT_PATH is not set in .env")
        sys.exit(1)

    today = datetime.date.today().isoformat()
    folder = Path(VAULT_PATH) / "Word_Summaries"
    folder.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading(f"Daily Learning Summary — {today}", level=1)

    for line in content.splitlines():
        stripped = line.strip()
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped:
            doc.add_paragraph(stripped)

    out = folder / f"{today}.docx"
    doc.save(str(out))
    print(f"Saved: {out}")


if __name__ == "__main__":
    # === RUTHLESS VAULT GUARD — refuse to operate on a diseased system ===
    try:
        from vault_guard import run_audit
        code = run_audit(strict=False)
        if code >= 2:
            print("\n[GUARD] CRITICAL: Vault failed health check. Fix issues before generating daily memory.")
            print("Your second brain is only as good as the discipline you enforce.")
            sys.exit(2)
    except Exception as e:
        print(f"[GUARD] Warning: could not run vault guard ({e}). Proceeding with reduced safety.")

    print("Extracting today's desktop memory context...")
    raw_logs = get_daily_text()
    today = datetime.date.today().isoformat()

    if raw_logs:
        print(f"Processing data ({len(raw_logs)} characters)... Generating AI summary.")
        summary = summarize_text(raw_logs)
        print("Exporting summary to Obsidian vault...")
        send_to_obsidian(summary)
        send_to_word(summary)

        print("Scanning for unfinished intentions in today's logs...")
        open_loops = extract_open_loops(raw_logs, today)
        if open_loops:
            append_to_open_loops(open_loops, today)
            print(f"  → {len(open_loops)} open loop(s) added to Memory/Open-Loops.md")
            print("  → Run: python vault_intelligence.py gaps --scope Memory  to rebuild the dashboard")
        else:
            print("  → No unfinished intentions detected in today's logs.")
    else:
        print("No active logs retrieved for today. Make sure Screenpipe is running!")
